# app.py (採用新輸出結構的最終版)

# --- Python 標準庫與第三方庫 ---
import os
import uuid
import json
import queue
import threading
import logging
import shutil
import locale
import sys
import configparser
import socket
import time

# --- Web 框架與 GUI ---
from flask import Flask, request, render_template, flash, redirect, url_for, Response, jsonify, session
from werkzeug.utils import secure_filename
import webview
import google.generativeai as genai
import docx

# --- 專案內部模組 ---
from ai_config import MODEL_CONFIG, PROMPTS
from workflow_scripts.pdf_ocr_translator import run_ocr_translation, run_ocr_only, run_ocr_translation_for_image
from workflow_scripts.text_summarizer import run_summarization
from workflow_scripts.summary_to_ppt import run_conversion_to_ppt
from workflow_scripts.pdf_splitter import run_pdf_split
from desktop_utils import copy_to_desktop_folder, open_folder_in_explorer

# ==============================================================================
#                                  應用程式設置
# ==============================================================================
def get_base_path():
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    else:
        return os.path.dirname(os.path.abspath(__file__))

BASE_PATH = get_base_path()
CONFIG_FILE = os.path.join(BASE_PATH, 'config.ini')
LOG_FILE = os.path.join(BASE_PATH, 'ai_toolkit.log')

logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s [%(levelname)s] %(threadName)s: %(message)s',
                    handlers=[
                        logging.FileHandler(LOG_FILE, mode='a', encoding='utf-8'),
                        logging.StreamHandler(sys.stdout)
                    ])

app = Flask(__name__)
UPLOAD_FOLDER = os.path.join(BASE_PATH, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_PATH, 'outputs')
ALLOWED_EXTENSIONS_PDF = {'pdf'}
ALLOWED_EXTENSIONS_OCR = {'pdf', 'png', 'jpg', 'jpeg', 'bmp', 'gif', 'webp'}
ALLOWED_EXTENSIONS_FULL_REPORT = {'pdf'}
# +++ 新增：允許的文字檔案類型 +++
ALLOWED_EXTENSIONS_TEXT = {'docx', 'txt'}


app.secret_key = os.environ.get("FLASK_SECRET_KEY", "a_very_secret_key_that_should_be_changed")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# ==============================================================================
#                              背景任務處理機制
# ==============================================================================
central_task_queue = queue.Queue()
task_progress_queues = {}

def task_worker():
    logging.info("[背景工作者] 工作者執行緒已啟動，等待任務...")
    while True:
        task_info = None
        task_id = None
        progress_queue = None
        try:
            task_info = central_task_queue.get()
            task_id = task_info.get('task_id')
            task_type = task_info.get('task_type')
            progress_queue = task_progress_queues.get(task_id)

            if not task_id or not task_type or not progress_queue:
                logging.error(f"[背景工作者] 從佇列收到無效的任務資訊或找不到進度佇列: {task_info}")
                central_task_queue.task_done()
                continue

            logging.info(f"[背景工作者] 開始處理任務 {task_id} (類型: {task_type})")
            api_key = app.config.get('GEMINI_API_KEY')
            if not api_key:
                raise ValueError(f"任務 {task_id} 缺少 GEMINI_API_KEY")
            
            genai.configure(api_key=api_key)

            if task_type == 'pdf_to_ppt':
                run_full_workflow(progress_queue, task_id, api_key, task_info)
            elif task_type == 'full_report':
                run_full_report_workflow(progress_queue, task_id, api_key, task_info)
            elif task_type == 'ocr':
                run_ocr_workflow(progress_queue, task_id, api_key, task_info)
            elif task_type == 'summarize':
                run_summarize_workflow(progress_queue, task_id, api_key, task_info)
            elif task_type == 'file_split':
                run_split_workflow(progress_queue, task_id, api_key, task_info)
            # +++ 新增：處理新的任務類型 +++
            elif task_type == 'text_to_ppt':
                run_text_to_ppt_workflow(progress_queue, task_id, api_key, task_info)
            else:
                logging.warning(f"[背景工作者] 未知的任務類型: {task_type} (ID: {task_id})")
                progress_queue.put(json.dumps({'type': 'error', 'message': f'未知的任務類型: {task_type}'}))

            logging.info(f"[背景工作者] 任務 {task_id} 處理完成。")

        except Exception as worker_e:
            logging.error(f"[背景工作者] 處理任務 {task_id} 時發生錯誤: {worker_e}", exc_info=True)
            if progress_queue:
                progress_queue.put(json.dumps({'type': 'error', 'message': f'處理任務時發生內部錯誤: {worker_e}'}))
        finally:
            central_task_queue.task_done()
            if task_id:
                task_progress_queues.pop(task_id, None)
                logging.debug(f"[背景工作者] 任務 {task_id} 已完成並清理其進度佇列。")

worker_thread = threading.Thread(target=task_worker, daemon=True, name="TaskWorkerThread")
worker_thread.start()

# ==============================================================================
#                                輔助函式
# ==============================================================================
def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def read_text_from_file(filepath):
    _, extension = os.path.splitext(filepath)
    extension = extension.lower()
    text = ""
    logging.info(f"嘗試讀取檔案: {filepath}")
    try:
        if extension == '.docx':
            doc = docx.Document(filepath)
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
        elif extension == '.txt':
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            except UnicodeDecodeError:
                default_encoding = locale.getpreferredencoding(False)
                with open(filepath, 'r', encoding=default_encoding) as f:
                    text = f.read()
        else:
            raise ValueError(f"不支援讀取文字的檔案類型: {extension}")
        logging.info(f"成功讀取檔案: {os.path.basename(filepath)} (文字長度: {len(text)})")
        return text
    except Exception as e:
        logging.error(f"讀取檔案失敗 ({filepath}): {e}", exc_info=True)
        raise

# ==============================================================================
#                                工作流程函式
# ==============================================================================

# +++ 新增：處理文字檔到簡報的工作流程 +++
def run_text_to_ppt_workflow(progress_queue, task_id, api_key, task_info):
    """從文字檔 (docx, txt) 開始，執行 摘要 -> 簡報 的流程。"""
    logging.info(f"[工作流程 {task_id} - 文字檔->PPT] 開始...")
    original_fn = task_info['original_base_filename_preserved']
    uploaded_file = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    
    summary_subfolder = "sum"
    ppt_subfolder = "ppt"
    
    summary_word_path = os.path.join(task_folder, "summary.docx")
    temp_ppt_path = os.path.join(task_folder, "ppt.pptx")
    
    overall_success = False
    output_path = None
    
    try:
        # 步驟 1: 讀取檔案內容
        progress_queue.put(json.dumps({'type': 'status', 'step': 1, 'status': '讀取檔案內容...', 'percent': 10}))
        document_text = read_text_from_file(uploaded_file)
        if not document_text.strip():
            raise ValueError("上傳的檔案內容為空或無法讀取。")

        # 步驟 2: 生成摘要
        progress_queue.put(json.dumps({'type': 'status', 'step': 2, 'status': '生成摘要...', 'percent': 25}))
        if not run_summarization(api_key, MODEL_CONFIG['SUMMARIZE'], document_text, summary_word_path, progress_queue):
            raise Exception("步驟 2 (生成摘要) 失敗")
        copy_to_desktop_folder(summary_word_path, summary_subfolder, f"sum_{original_fn}.docx")

        # 步驟 3: 生成簡報
        progress_queue.put(json.dumps({'type': 'status', 'step': 3, 'status': '生成簡報...', 'percent': 80}))
        if not run_conversion_to_ppt(summary_word_path, temp_ppt_path):
            raise Exception("步驟 3 (轉換為簡報) 失敗")

        final_path, final_ppt_name = copy_to_desktop_folder(temp_ppt_path, ppt_subfolder, f"ppt_{original_fn}.pptx")
        if not final_path:
            raise Exception("儲存最終簡報到桌面失敗")
        
        output_path = os.path.dirname(os.path.dirname(final_path))
        progress_queue.put(json.dumps({
            'type': 'complete', 
            'message': f'簡報 "{final_ppt_name}" 及摘要檔案已分類儲存。',
            'folder_path': output_path
        }))
        overall_success = True

    except Exception as e:
        logging.error(f"[工作流程 {task_id} - 文字檔->PPT] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder):
            shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success:
            progress_queue.put(json.dumps({'type': 'done'}))


def run_full_report_workflow(progress_queue, task_id, api_key, task_info):
    """執行 PDF -> 原文 -> 翻譯 -> 摘要 -> 簡報 的完整流程，並將4個檔案分類儲存。"""
    logging.info(f"[工作流程 {task_id} - 完整簡報生成] 開始...")
    original_fn = task_info['original_base_filename_preserved']
    uploaded_pdf = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    
    ocr_subfolder = "ocr"
    trans_subfolder = "trans"
    summary_subfolder = "sum"
    ppt_subfolder = "ppt"
    
    ocr_path = os.path.join(task_folder, "ocr.docx")
    trans_path = os.path.join(task_folder, "translated.docx")
    summary_path = os.path.join(task_folder, "summary.docx")
    ppt_path = os.path.join(task_folder, "ppt.pptx")
    
    overall_success = False
    output_path = None
    try:
        # 步驟 1: 純 OCR 掃描原文
        progress_queue.put(json.dumps({'type': 'status', 'step': 1, 'status': '掃描原文 (OCR)...', 'percent': 5}))
        if not run_ocr_only(api_key, MODEL_CONFIG['OCR'], uploaded_pdf, ocr_path, progress_queue):
            raise Exception("步驟 1 (掃描原文) 失敗")
        copy_to_desktop_folder(ocr_path, ocr_subfolder, f"ocr_{original_fn}.docx")
        
        # 讀取 OCR 結果以進行下一步翻譯
        ocr_text = read_text_from_file(ocr_path)
        if not ocr_text.strip(): raise ValueError("掃描後的原文內容為空")

        # 步驟 2: 翻譯原文
        progress_queue.put(json.dumps({'type': 'status', 'step': 2, 'status': '翻譯原文...', 'percent': 30}))
        translation_prompt = f"請將以下全文精確地翻譯成繁體中文，並盡可能保持原有的格式和段落結構。請不要添加任何摘要或評論，只需純粹的翻譯。\n\n---\n{ocr_text}\n---"
        # 使用成本較低的模型進行純文字翻譯
        trans_model = genai.GenerativeModel(MODEL_CONFIG['OCR']) 
        response = trans_model.generate_content(translation_prompt)
        if not hasattr(response, 'text') or not response.text: 
            raise Exception("步驟 2 (翻譯) 失敗: AI 未返回有效的翻譯內容。")
        
        translated_text = response.text
        trans_doc = docx.Document()
        trans_doc.add_paragraph(translated_text)
        trans_doc.save(trans_path)
        copy_to_desktop_folder(trans_path, trans_subfolder, f"trans_{original_fn}.docx")

        # 步驟 3: 生成摘要
        progress_queue.put(json.dumps({'type': 'status', 'step': 3, 'status': '生成摘要...', 'percent': 55}))
        if not run_summarization(api_key, MODEL_CONFIG['SUMMARIZE'], translated_text, summary_path, progress_queue):
            raise Exception("步驟 3 (生成摘要) 失敗")
        copy_to_desktop_folder(summary_path, summary_subfolder, f"sum_{original_fn}.docx")

        # 步驟 4: 生成簡報
        progress_queue.put(json.dumps({'type': 'status', 'step': 4, 'status': '生成簡報...', 'percent': 80}))
        if not run_conversion_to_ppt(summary_path, ppt_path):
            raise Exception("步驟 4 (轉換為簡報) 失敗")
        
        final_path, _ = copy_to_desktop_folder(ppt_path, ppt_subfolder, f"ppt_{original_fn}.pptx")
        if not final_path: raise Exception("儲存最終簡報到桌面失敗")
        
        output_path = os.path.dirname(os.path.dirname(final_path))
        progress_queue.put(json.dumps({
            'type': 'complete', 
            'message': '完整報告處理完成，4類檔案已分類儲存。',
            'folder_path': output_path
        }))
        overall_success = True

    except Exception as e:
        logging.error(f"[工作流程 {task_id} - 完整簡報生成] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder): shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success: progress_queue.put(json.dumps({'type': 'done'}))


def run_full_workflow(progress_queue, task_id, api_key, task_info):
    logging.info(f"[工作流程 {task_id} - PDF->PPT] 開始...")
    original_fn = task_info['original_base_filename_preserved']
    uploaded_pdf = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    
    trans_subfolder = "trans"
    summary_subfolder = "sum"
    ppt_subfolder = "ppt"
    
    step1_word_path = os.path.join(task_folder, "translated.docx")
    summary_word_path = os.path.join(task_folder, "summary.docx")
    temp_ppt_path = os.path.join(task_folder, "ppt.pptx")
    overall_success = False
    output_path = None
    try:
        progress_queue.put(json.dumps({'type': 'status', 'step': 1, 'status': 'OCR與翻譯...', 'percent': 5}))
        if not run_ocr_translation(api_key, MODEL_CONFIG['OCR'], uploaded_pdf, step1_word_path, progress_queue):
            raise Exception("步驟 1 (OCR/翻譯) 失敗")
        copy_to_desktop_folder(step1_word_path, trans_subfolder, f"trans_{original_fn}.docx")

        progress_queue.put(json.dumps({'type': 'status', 'step': 2, 'status': '生成摘要...', 'percent': 40}))
        step1_text = read_text_from_file(step1_word_path)
        if not step1_text.strip(): raise ValueError("翻譯檔案內容為空")
        if not run_summarization(api_key, MODEL_CONFIG['SUMMARIZE'], step1_text, summary_word_path, progress_queue):
            raise Exception("步驟 2 (生成摘要) 失敗")
        copy_to_desktop_folder(summary_word_path, summary_subfolder, f"sum_{original_fn}.docx")
        
        progress_queue.put(json.dumps({'type': 'status', 'step': 3, 'status': '生成簡報...', 'percent': 80}))
        if not run_conversion_to_ppt(summary_word_path, temp_ppt_path):
            raise Exception("步驟 3 (轉換為簡報) 失敗")

        final_path, final_ppt_name = copy_to_desktop_folder(temp_ppt_path, ppt_subfolder, f"ppt_{original_fn}.pptx")
        if not final_path: raise Exception("儲存最終簡報到桌面失敗")
        
        output_path = os.path.dirname(os.path.dirname(final_path))
        progress_queue.put(json.dumps({
            'type': 'complete', 
            'message': f'快速簡報 "{final_ppt_name}" 及過程檔案已分類儲存。',
            'folder_path': output_path
        }))
        overall_success = True

    except Exception as e:
        logging.error(f"[工作流程 {task_id} - PDF->PPT] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder): shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success: progress_queue.put(json.dumps({'type': 'done'}))


def run_ocr_workflow(progress_queue, task_id, api_key, task_info):
    logging.info(f"[工作流程 {task_id} - OCR] 開始...")
    original_fn = task_info['original_base_filename_preserved']
    uploaded_file = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    temp_docx_path = os.path.join(task_folder, "ocr_result.docx")
    output_subfolder = "trans"
    overall_success = False
    output_path = None
    try:
        progress_queue.put(json.dumps({'type': 'status', 'step': 1, 'status': 'OCR 與翻譯處理中...', 'percent': 5}))
        _, ext = os.path.splitext(uploaded_file)
        if ext.lower() == '.pdf':
            success = run_ocr_translation(api_key, MODEL_CONFIG['OCR'], uploaded_file, temp_docx_path, progress_queue)
        else:
            success = run_ocr_translation_for_image(api_key, MODEL_CONFIG['OCR'], uploaded_file, temp_docx_path, progress_queue)
        if not success: raise Exception("OCR 與翻譯步驟失敗")

        final_path, final_name = copy_to_desktop_folder(temp_docx_path, output_subfolder, f"trans_{original_fn}.docx")
        if not final_path: raise Exception("儲存檔案到桌面失敗")

        output_path = os.path.dirname(os.path.dirname(final_path))
        progress_queue.put(json.dumps({
            'type': 'complete',
            'message': f'翻譯檔案 "{final_name}" 已儲存。',
            'folder_path': output_path
        }))
        overall_success = True
    except Exception as e:
        logging.error(f"[工作流程 {task_id} - OCR] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder): shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success: progress_queue.put(json.dumps({'type': 'done'}))


def run_summarize_workflow(progress_queue, task_id, api_key, task_info):
    logging.info(f"[工作流程 {task_id} - Summarize] 開始...")
    original_fn = task_info['original_base_filename_preserved']
    uploaded_pdf = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    ocr_word_path = os.path.join(task_folder, "ocr_output.docx")
    temp_summary_path = os.path.join(task_folder, "summary.docx")
    output_subfolder = "sum"
    overall_success = False
    output_path = None
    try:
        if not run_ocr_translation(api_key, MODEL_CONFIG['OCR'], uploaded_pdf, ocr_word_path, progress_queue):
            raise Exception("步驟 1 (OCR) 失敗")
        
        doc_text = read_text_from_file(ocr_word_path)
        if not doc_text.strip(): raise ValueError("OCR 結果為空")
        if not run_summarization(api_key, MODEL_CONFIG['SUMMARIZE'], doc_text, temp_summary_path, progress_queue):
            raise Exception("步驟 2 (摘要) 失敗")
        
        final_path, final_name = copy_to_desktop_folder(temp_summary_path, output_subfolder, f"sum_{original_fn}.docx")
        if not final_path: raise Exception("儲存檔案到桌面失敗")

        output_path = os.path.dirname(os.path.dirname(final_path))
        progress_queue.put(json.dumps({
            'type': 'complete',
            'message': f'摘要檔案 "{final_name}" 已儲存。',
            'folder_path': output_path
        }))
        overall_success = True
    except Exception as e:
        logging.error(f"[工作流程 {task_id} - Summarize] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder): shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success: progress_queue.put(json.dumps({'type': 'done'}))


def run_split_workflow(progress_queue, task_id, api_key, task_info):
    logging.info(f"[工作流程 {task_id} - FileSplit] 開始...")
    uploaded_pdf = task_info['uploaded_file_path']
    task_folder = task_info['task_output_folder']
    output_subfolder = "檔案分割輸出"
    overall_success = False
    try:
        split_count, output_dir = run_pdf_split(
            api_key=api_key,
            model_name=MODEL_CONFIG['PDF_SPLIT_ANALYSIS'],
            input_pdf_path=uploaded_pdf,
            output_folder_name=output_subfolder,
            progress_queue=progress_queue
        )
        if split_count > 0:
            progress_queue.put(json.dumps({
                'type': 'complete',
                'message': f'成功分割成 {split_count} 個檔案。',
                'folder_path': output_dir
            }))
            overall_success = True
        elif split_count == 0:
             raise Exception("AI 未能分析出有效的目錄結構，無法分割。")
        else:
            raise Exception("檔案分割過程中發生未知錯誤。")
    except Exception as e:
        logging.error(f"[工作流程 {task_id} - FileSplit] 失敗: {e}", exc_info=True)
        progress_queue.put(json.dumps({'type': 'error', 'message': f'處理失敗: {e}'}))
    finally:
        if os.path.exists(task_folder): shutil.rmtree(task_folder, ignore_errors=True)
        if not overall_success: progress_queue.put(json.dumps({'type': 'done'}))

# ==============================================================================
#                                Flask 路由
# ==============================================================================
@app.route('/')
def index():
    return render_template('landing.html')

@app.route('/pdf_to_ppt')
def pdf_to_ppt():
    page_context = {
        "title": "生成簡報 (快速)",
        "icon": "bi-camera-reels",
        "description": "選擇 PDF，系統將自動 OCR、翻譯、摘要並生成 PPTX 簡報。過程檔案與最終簡報將分類儲存到桌面。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ".pdf",
        "task_type": "pdf_to_ppt",
        "button_text": "開始生成",
        "button_color_class": "btn-grad-1",
        "output_folder_name": "AI 工具輸出"
    }
    return render_template('process_page.html', **page_context)

@app.route('/full_report')
def full_report():
    page_context = {
        "title": "完整簡報生成 (全流程)",
        "icon": "bi-journal-album",
        "description": "這是一個完整的處理流程。選擇 PDF 後，系統將依序產出【原文掃描檔】、【翻譯檔】、【摘要檔】和最終的【簡報檔】，並分類儲存。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ".pdf",
        "task_type": "full_report",
        "button_text": "開始完整生成",
        "button_color_class": "btn-grad-special",
        "output_folder_name": "AI 工具輸出"
    }
    return render_template('process_page.html', **page_context)

@app.route('/ocr')
def ocr():
    page_context = {
        "title": "文字辨識 (含翻譯)",
        "icon": "bi-textarea-t",
        "description": "選擇 PDF 或圖片檔案，系統將進行 OCR 並翻譯為繁體中文。結果將儲存到桌面的「翻譯檔案」資料夾。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ", ".join(f".{ext}" for ext in ALLOWED_EXTENSIONS_OCR),
        "task_type": "ocr",
        "button_text": "開始辨識翻譯",
        "button_color_class": "btn-grad-2",
        "output_folder_name": "AI 工具輸出/trans"
    }
    return render_template('process_page.html', **page_context)

@app.route('/summarize')
def summarize():
    page_context = {
        "title": "重點整理",
        "icon": "bi-card-checklist",
        "description": "選擇 PDF 文件，系統將進行 OCR、翻譯並生成結構化的重點摘要。結果將儲存到桌面的「摘要檔案」資料夾。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ".pdf",
        "task_type": "summarize",
        "button_text": "開始整理",
        "button_color_class": "btn-grad-3",
        "output_folder_name": "AI 工具輸出/sum"
    }
    return render_template('process_page.html', **page_context)

# +++ 新增：文字檔到簡報的路由 +++
@app.route('/text_to_ppt')
def text_to_ppt():
    page_context = {
        "title": "文字檔生成簡報",
        "icon": "bi-file-earmark-text",
        "description": "上傳 Word (.docx) 或純文字檔 (.txt)，系統將直接為其生成結構化摘要與 PowerPoint 簡報。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ".docx, .txt",
        "task_type": "text_to_ppt",
        "button_text": "開始生成",
        "button_color_class": "btn-grad-6", # 使用一個新的顏色
        "output_folder_name": "AI 工具輸出"
    }
    return render_template('process_page.html', **page_context)


@app.route('/file_split')
def file_split():
    page_context = {
        "title": "PDF 智能分割",
        "icon": "bi-scissors",
        "description": "選擇 PDF 文件，系統將使用 AI 分析其目錄結構，並按章節將其分割成多個獨立的 PDF 檔案儲存到桌面。",
        "form_action_url": url_for('process_task'),
        "allowed_extensions": ".pdf",
        "task_type": "file_split",
        "button_text": "開始分割",
        "button_color_class": "btn-grad-5",
        "output_folder_name": "AI 工具輸出/檔案分割輸出/[檔名]"
    }
    return render_template('process_page.html', **page_context)

@app.route('/process_task', methods=['POST'])
def process_task():
    task_type = request.form.get('task_type')
    if not task_type: return jsonify({'success': False, 'error': '未知的任務類型'}), 400
    if 'source_file' not in request.files: return jsonify({'success': False, 'error': '沒有選擇檔案'}), 400
    file = request.files['source_file']
    if file.filename == '': return jsonify({'success': False, 'error': '沒有選擇檔案'}), 400
    
    # +++ 更新：加入新的任務類型和允許的副檔名 +++
    allowed_ext_map = {
        'pdf_to_ppt': ALLOWED_EXTENSIONS_PDF,
        'full_report': ALLOWED_EXTENSIONS_FULL_REPORT,
        'ocr': ALLOWED_EXTENSIONS_OCR,
        'summarize': ALLOWED_EXTENSIONS_PDF,
        'file_split': ALLOWED_EXTENSIONS_PDF,
        'text_to_ppt': ALLOWED_EXTENSIONS_TEXT,
    }
    allowed_exts = allowed_ext_map.get(task_type)
    
    if not allowed_exts: return jsonify({'success': False, 'error': f'不支援的任務類型: {task_type}'}), 400
    if not allowed_file(file.filename, allowed_exts): return jsonify({'success': False, 'error': f'檔案類型不支援，請上傳 {"/".join(allowed_exts)} 檔案'}), 400
    
    original_full_filename = file.filename
    original_base = os.path.splitext(original_full_filename)[0]
    safe_filename = secure_filename(original_full_filename)

    task_id = str(uuid.uuid4())
    task_output_folder = os.path.join(app.config['OUTPUT_FOLDER'], task_id)
    os.makedirs(task_output_folder, exist_ok=True)
    uploaded_file_path = os.path.join(task_output_folder, safe_filename)
    try:
        file.save(uploaded_file_path)
    except Exception as e:
        shutil.rmtree(task_output_folder, ignore_errors=True)
        return jsonify({'success': False, 'error': f'儲存上傳檔案失敗: {e}'}), 500

    task_progress_queues[task_id] = queue.Queue()
    task_info = {'task_id': task_id, 'task_type': task_type, 'original_base_filename_preserved': original_base, 'uploaded_file_path': uploaded_file_path, 'task_output_folder': task_output_folder}
    central_task_queue.put(task_info)
    return jsonify({'success': True, 'task_id': task_id, 'filename': original_full_filename})

@app.route('/stream/<task_id>')
def stream(task_id):
    def sse_event_stream():
        logging.info(f"[SSE {task_id}] 客戶端已連接")
        q = task_progress_queues.get(task_id)
        if q is None: yield f"data: {json.dumps({'type':'error', 'message':'任務已完成或不存在。'})}\n\n"; return
        done = False
        try:
            while not done:
                try:
                    data_str = q.get(timeout=60)
                    data = json.loads(data_str)
                    yield f"data: {data_str}\n\n"
                    if data.get('type') in ['done', 'error', 'complete']: done = True
                except queue.Empty: yield ":keep-alive\n\n"
                except (json.JSONDecodeError, TypeError): logging.warning(f"[SSE {task_id}] 從佇列收到無效資料: {data_str if 'data_str' in locals() else 'N/A'}"); continue
        except GeneratorExit: logging.info(f"[SSE {task_id}] 客戶端已斷開連接")
        finally: logging.info(f"[SSE {task_id}] 事件串流結束。")
    try: uuid.UUID(task_id)
    except ValueError: return Response("Invalid task ID format", status=400)
    return Response(sse_event_stream(), mimetype="text/event-stream")

@app.route('/chat')
def chat():
    session.pop('chat_history', None)
    return render_template('chat.html')

@app.route('/api/chat', methods=['POST'])
def api_chat():
    api_key = app.config.get('GEMINI_API_KEY')
    if not api_key: return jsonify({'error': 'API Key not configured'}), 500
    data = request.json
    user_message = data.get('message')
    if not user_message: return jsonify({'error': 'No message provided'}), 400
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(MODEL_CONFIG['CHAT'])
        chat_history = session.get('chat_history', [])
        temp_chat = model.start_chat(history=chat_history)
        response = temp_chat.send_message(user_message)
        ai_reply = response.text
        chat_history.append({'role': 'user', 'parts': [user_message]})
        chat_history.append({'role': 'model', 'parts': [ai_reply]})
        session['chat_history'] = chat_history
        return jsonify({'reply': ai_reply})
    except Exception as e:
        logging.error(f"呼叫 Chat API 時發生錯誤: {e}", exc_info=True)
        return jsonify({'error': '與 AI 溝通時發生內部錯誤。'}), 500

@app.route('/open_folder', methods=['POST'])
def open_folder():
    data = request.json
    folder_path = data.get('path')
    if not folder_path or not os.path.isdir(folder_path): return jsonify({'success': False, 'error': '無效的路徑'}), 400
    try: open_folder_in_explorer(folder_path); return jsonify({'success': True})
    except Exception as e: logging.error(f"從 API 端點打開資料夾失敗: {e}"); return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/translate')
def translate():
    flash('「即時翻譯」功能仍在開發中，敬請期待！', 'info')
    return redirect(url_for('index'))

# ==============================================================================
#                      桌面應用程式啟動器
# ==============================================================================
if __name__ == '__main__':
    logging.info("============================================"); logging.info("====== AI 工具箱桌面應用程式啟動中... ====="); logging.info(f"    基礎路徑: {BASE_PATH}"); logging.info("============================================")
    def get_api_key_from_config():
        if not os.path.exists(CONFIG_FILE):
            logging.error(f"設定檔 {CONFIG_FILE} 不存在。")
            try:
                with open(CONFIG_FILE, 'w', encoding='utf-8') as f: f.write("[Credentials]\n"); f.write("GEMINI_API_KEY = YOUR_GEMINI_API_KEY_HERE\n")
                return None, f"設定檔不存在。已建立模板 {os.path.basename(CONFIG_FILE)}，請填入您的 API Key 後重啟。"
            except Exception as e_cfg: return None, f"設定檔不存在且無法自動建立: {e_cfg}"
        config = configparser.ConfigParser()
        try:
            config.read(CONFIG_FILE, encoding='utf-8')
            api_key = config.get('Credentials', 'GEMINI_API_KEY', fallback="").strip()
            if not api_key or api_key == 'YOUR_GEMINI_API_KEY_HERE': return None, f"請在設定檔 {os.path.basename(CONFIG_FILE)} 中提供有效的 GEMINI_API_KEY。"
            return api_key, None
        except Exception as e: return None, f"讀取設定檔時出錯: {e}"
    loaded_api_key, config_error_msg = get_api_key_from_config()
    if not loaded_api_key:
        logging.critical(f"因設定錯誤無法啟動: {config_error_msg}")
        webview.create_window("設定錯誤", html=f"<h1>設定錯誤</h1><p>{config_error_msg}</p>", width=500, height=200); webview.start(); sys.exit(1)
    app.config['GEMINI_API_KEY'] = loaded_api_key; logging.info("GEMINI_API_KEY 已載入到 Flask 設定。")
    def find_free_port():
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s: s.bind(("127.0.0.1", 0)); return s.getsockname()[1]
    port = find_free_port()
    server_thread = threading.Thread(target=lambda: app.run(host="127.0.0.1", port=port, debug=False, use_reloader=False)); server_thread.daemon = True; server_thread.start()
    logging.info(f"Flask 伺服器執行緒已在 http://127.0.0.1:{port} 啟動")
    try:
        webview.create_window("AI 工具庫", f"http://127.0.0.1:{port}/", width=1100, height=750, resizable=True, confirm_close=True)
        webview.start(debug=False)
    except Exception as e_webview: logging.critical(f"建立 pywebview 視窗失敗: {e_webview}", exc_info=True); sys.exit(1)
    logging.info("pywebview 視窗已關閉，應用程式結束。"); sys.exit(0)