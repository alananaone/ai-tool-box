# workflow_scripts/text_summarizer.py (修改版)
import google.generativeai as genai
import os
import docx
import time
import re
import json # 用於進度回報
import logging # 使用 logging

# --- Prompt 保持不變 ---
SUMMARY_PROMPT = """
請仔細閱讀並理解以下提供的文件全文。
你的任務是為這份文件生成一份【詳盡】、【保留所有重要資訊】且【結構清晰、層次分明】的【繁體中文】摘要。

請嚴格遵循以下指示進行摘要與格式化：

1.  **識別主要主題：** 從文本中識別出最高層級的主要主題或章節。
2.  **識別子主題：** 在每個主要主題下，識別出次要的子主題或代表原始文件主要段落的核心概念。
3.  **使用 Markdown 標示層級與細節：**
    * 對於識別出的**主要主題/章節**，使用 Markdown 的 Heading 1 語法 (`# 主要主題文字`)。
    * 在每個主要主題下方，對於**子主題/核心概念**，使用 Markdown 的 Heading 2 語法 (`## 子主題文字`)。
    * 在每個 `## 子主題` 下方，使用**項目符號 (bullet points)** (`* ` 或 `- `) 來條列式摘要該子主題相關的所有關鍵定義、細節、範例、效果、風險等。
        * **重要：** 這些項目符號的內容和順序，應盡可能**反映原文件中與該子主題相關段落的資訊流和結構**。
        * **重要：** 如果某個要點下有更細節的說明或從屬資訊，請在提示模型時想像使用**縮排**來表示下一層的項目符號（雖然輸出仍主要依賴 `* `，但要求模型思考層次關係）。腳本目前主要處理第一層 `* ` 轉換為 Word 的項目符號。
4.  **內容要求：**
    * 摘要必須【極力保留原文的所有關鍵資訊】，包括但不限於定義、分類、數據、背景、機制、影響、用途、問題、風險、結論等。**確保不遺漏重要細節**。
    * 確保摘要的邏輯流暢，且項目符號下的內容確實反映了原文相應部分的重點。
5.  **格式要求：** 確保輸出的 Markdown 格式（主要依賴 `#`, `##`, `* `）正確、層次清晰，以便後續程式能準確解析。
6.  **輸出限制：** 請【只輸出】符合上述格式要求的完整 Markdown 摘要內容，不要包含任何額外的前言（例如 "這是您的摘要："）、結語或與摘要無關的文字。輸出應直接從第一個 `# 主要主題` 開始。

以下是需要摘要的文件全文：
---
{document_text}
---
"""
# --- Prompt 結束 ---

API_DELAY = 1 # 秒

# --- 移除 get_text_from_docx，因為文字會在 app.py 中讀取 ---
# def get_text_from_docx(filepath): ...

# +++ 修改函式簽名：接收 document_text 和 progress_queue +++
def run_summarization(api_key: str, model_name: str, document_text: str, output_summary_path: str, progress_queue=None) -> bool:
    """
    根據提供的文字內容生成摘要，並將結果儲存為 Word 文件。

    Args:
        api_key: Gemini API 金鑰。
        model_name: 要使用的 Gemini 模型名稱。
        document_text: 要摘要的完整文字內容。
        output_summary_path: 輸出摘要 Word 檔案的路徑。
        progress_queue: 用於傳遞進度訊息的 queue.Queue 物件 (可選)。

    Returns:
        bool: 成功時返回 True，失敗時返回 False。
    """
    logging.info(f"開始生成摘要...")
    if progress_queue:
        progress_queue.put(json.dumps({'type': 'status', 'status': '初始化摘要模型...', 'percent': 10})) # 提供初始進度

    try:
        # 假設 API Key 已在外部配置
        # genai.configure(api_key=api_key)
        model = genai.GenerativeModel(model_name)
        logging.info(f"  使用的摘要模型: {model_name}")
    except Exception as e:
        logging.error(f"  設定 Gemini 或建立模型時發生錯誤: {e}")
        if progress_queue:
            progress_queue.put(json.dumps({'type': 'error', 'message': f'建立摘要模型失敗: {e}'}))
        return False

    # 檢查輸入文字是否有效
    if not document_text or not document_text.strip():
        logging.warning("  輸入的文字內容為空。")
        if progress_queue:
            progress_queue.put(json.dumps({'type': 'error', 'message': '輸入的文字內容為空。'}))
        return False

    logging.info("  文字內容有效，呼叫 Gemini API 生成摘要...")
    if progress_queue:
        progress_queue.put(json.dumps({'type': 'status', 'status': '正在呼叫 AI 生成摘要...', 'percent': 30}))

    try:
        # 使用傳入的 document_text 格式化 Prompt
        prompt_with_text = SUMMARY_PROMPT.format(document_text=document_text)
        response = model.generate_content(prompt_with_text)

        # API 延遲
        time.sleep(API_DELAY)

        if hasattr(response, 'text') and response.text:
            summary_markdown = response.text.strip()
            logging.info("  摘要生成成功，正在寫入 Word 檔案...")
            if progress_queue:
                progress_queue.put(json.dumps({'type': 'status', 'status': '正在格式化並儲存摘要檔案...', 'percent': 80}))

            try:
                summary_doc = docx.Document()
                lines = summary_markdown.splitlines()
                for line in lines:
                    line_stripped = line.strip()
                    if not line_stripped: continue

                    # --- Word 文件寫入邏輯保持不變 ---
                    if line_stripped.startswith('# '):
                        text = re.sub(r'^#\s+', '', line_stripped)
                        para = summary_doc.add_paragraph(text)
                        para.style = 'Heading 1'
                    elif line_stripped.startswith('## '):
                        text = re.sub(r'^##\s+', '', line_stripped)
                        para = summary_doc.add_paragraph(text)
                        para.style = 'Heading 2'
                    elif line_stripped.startswith('### '):
                         text = re.sub(r'^###\s+', '', line_stripped)
                         para = summary_doc.add_paragraph(text)
                         para.style = 'Heading 3'
                    elif line_stripped.startswith('* ') or line_stripped.startswith('- '):
                         text = re.sub(r'^[*\-]\s+', '', line_stripped)
                         para = summary_doc.add_paragraph(text)
                         para.style = 'List Bullet'
                    else:
                        para = summary_doc.add_paragraph(line_stripped)
                # --- Word 文件寫入邏輯結束 ---

                summary_doc.save(output_summary_path)
                logging.info(f"  摘要 Word 檔案儲存成功: {os.path.basename(output_summary_path)}")
                # 成功訊息由 workflow 函式發送
                return True

            except Exception as write_e:
                logging.error(f"  !! 錯誤: 寫入摘要 Word 檔案 '{os.path.basename(output_summary_path)}' 時失敗: {write_e}")
                if progress_queue:
                    progress_queue.put(json.dumps({'type': 'error', 'message': f'寫入摘要檔案失敗: {write_e}'}))
                return False
        else:
            block_reason = ""
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 block_reason = f" (Block Reason: {response.prompt_feedback.block_reason})"
            logging.warning(f"  !! 警告: Gemini API 未能生成有效的摘要文字{block_reason}。")
            if progress_queue:
                 progress_queue.put(json.dumps({'type': 'error', 'message': f'AI 未能生成摘要{block_reason}'}))
            return False

    except Exception as api_e:
        logging.error(f"  !! 錯誤: 呼叫 Gemini API 時發生錯誤: {api_e}", exc_info=True)
        if progress_queue:
             progress_queue.put(json.dumps({'type': 'error', 'message': f'呼叫 AI 時發生錯誤: {api_e}'}))
        return False

# --- (可以保留 if __name__ == '__main__': 用於單獨測試) ---
# if __name__ == '__main__':
#     # 測試代碼需要提供 API Key, 模型名稱, 測試文字, 和輸出路徑
#     logging.basicConfig(level=logging.INFO)
#     api_key = "YOUR_API_KEY"
#     model = "gemini-1.5-flash-latest"
#     test_text = "這是一段需要摘要的測試文字。\n它包含多個段落。\n目的是測試 run_summarization 函式。"
#     output_path = "test_summary_output.docx"
#     test_q = queue.Queue() # 模擬進度佇列
#
#     if not api_key or api_key == "YOUR_API_KEY":
#         print("請在腳本中設定您的 API Key")
#     else:
#         success = run_summarization(api_key, model, test_text, output_path, test_q)
#         print(f"測試執行結果: {'成功' if success else '失敗'}")
#         while not test_q.empty():
#             print("Queue message:", test_q.get())
#         if success and os.path.exists(output_path):
#             print(f"測試輸出檔案已生成: {output_path}")
#             # os.remove(output_path) # 可選：測試後刪除檔案
