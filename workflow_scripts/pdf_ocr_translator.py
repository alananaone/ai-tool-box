# workflow_scripts/pdf_ocr_translator.py (修正版)
import google.generativeai as genai
import os
import fitz  # PyMuPDF
from PIL import Image
import docx
import time
import json
import logging
from ai_config import PROMPTS # <--- 從 ai_config 導入 Prompts

API_DELAY = 0.5

def _process_pdf_pages(model, prompt_text, input_pdf_path, progress_queue):
    """通用內部函式，用於處理 PDF 頁面並返回 AI 生成的文字。"""
    full_text = ""
    page_errors = 0
    pdf_document = fitz.open(input_pdf_path)
    num_pages = len(pdf_document)
    logging.info(f"  PDF 共有 {num_pages} 頁，使用 Prompt: '{prompt_text[:20]}...'")
    if progress_queue:
        progress_queue.put(json.dumps({'type': 'progress', 'current': 0, 'total': num_pages, 'status': '開始處理頁面...'}))

    for page_num in range(num_pages):
        current_page_for_report = page_num + 1
        logging.info(f"    處理第 {current_page_for_report}/{num_pages} 頁...")
        if progress_queue:
            progress_queue.put(json.dumps({'type': 'progress', 'current': current_page_for_report, 'total': num_pages, 'status': f'處理中... ({current_page_for_report}/{num_pages})'}))

        page = pdf_document.load_page(page_num)
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
            img_bytes = pix.tobytes("png")
            image_part = {"mime_type": "image/png", "data": img_bytes}
            
            response = model.generate_content([prompt_text, image_part])

            if hasattr(response, 'text') and response.text:
                page_text = response.text.strip()
                full_text += page_text + "\n\n"
            else:
                page_errors += 1
                block_reason = f" (Block Reason: {response.prompt_feedback.block_reason})" if hasattr(response, 'prompt_feedback') else ""
                logging.warning(f"    頁面 {current_page_for_report}: [警告: 未生成文字{block_reason}]")
                full_text += f"[--- 第 {current_page_for_report} 頁處理失敗{block_reason} ---]\n\n"
            
            time.sleep(API_DELAY)
        except Exception as page_e:
            page_errors += 1
            logging.error(f"    頁面 {current_page_for_report}: [錯誤: {page_e}]")
            full_text += f"[--- 第 {current_page_for_report} 頁處理錯誤: {page_e} ---]\n\n"
            time.sleep(API_DELAY)
    
    pdf_document.close()
    if page_errors > 0:
        logging.warning(f"  注意：處理過程中出現 {page_errors} 個頁面錯誤。")
    
    return full_text

def run_ocr_translation(api_key: str, model_name: str, input_pdf_path: str, output_word_path: str, progress_queue=None) -> bool:
    """對 PDF 執行 OCR 和翻譯，結果儲存為 Word。"""
    logging.info(f"開始 OCR 與翻譯: {os.path.basename(input_pdf_path)}")
    try:
        model = genai.GenerativeModel(model_name)
        translated_text = _process_pdf_pages(model, PROMPTS["OCR_TRANSLATE"], input_pdf_path, progress_queue)
        
        if translated_text.strip():
            doc = docx.Document()
            doc.add_paragraph(translated_text)
            doc.save(output_word_path)
            logging.info(f"  翻譯 Word 檔案儲存成功: {os.path.basename(output_word_path)}")
            return True
        else:
            logging.warning("!! 警告: 未能從此 PDF 檔案中取得任何翻譯文字。")
            if progress_queue: progress_queue.put(json.dumps({'type': 'error', 'message': '未能取得任何翻譯文字'}))
            return False
    except Exception as e:
        logging.error(f"!! 嚴重錯誤: 執行 OCR 與翻譯時發生錯誤: {e}", exc_info=True)
        if progress_queue: progress_queue.put(json.dumps({'type': 'error', 'message': f'處理 PDF 檔案失敗: {e}'}))
        return False

def run_ocr_only(api_key: str, model_name: str, input_pdf_path: str, output_word_path: str, progress_queue=None) -> bool:
    """【只】對 PDF 執行 OCR (不翻譯)，結果儲存為 Word。"""
    logging.info(f"開始僅 OCR: {os.path.basename(input_pdf_path)}")
    try:
        model = genai.GenerativeModel(model_name)
        ocr_text = _process_pdf_pages(model, PROMPTS["OCR_ONLY"], input_pdf_path, progress_queue)

        if ocr_text.strip():
            doc = docx.Document()
            doc.add_paragraph(ocr_text)
            doc.save(output_word_path)
            logging.info(f"  OCR Word 檔案儲存成功: {os.path.basename(output_word_path)}")
            return True
        else:
            logging.warning("!! 警告: 未能從此 PDF 檔案中取得任何 OCR 文字。")
            if progress_queue: progress_queue.put(json.dumps({'type': 'error', 'message': '未能取得任何 OCR 文字'}))
            return False
    except Exception as e:
        logging.error(f"!! 嚴重錯誤: 執行僅 OCR 時發生錯誤: {e}", exc_info=True)
        if progress_queue: progress_queue.put(json.dumps({'type': 'error', 'message': f'處理 PDF 檔案失敗: {e}'}))
        return False

def run_ocr_translation_for_image(api_key: str, model_name: str, input_image_path: str, output_docx_path: str, progress_queue=None) -> bool:
    """對單張圖片執行 OCR 和翻譯。"""
    logging.info(f"開始處理圖片 OCR 與翻譯: {os.path.basename(input_image_path)}")
    try:
        model = genai.GenerativeModel(model_name)
        img = Image.open(input_image_path)
        
        if progress_queue: progress_queue.put(json.dumps({'type': 'status', 'status': '呼叫 AI 進行辨識翻譯...'}))
        
        response = model.generate_content([PROMPTS["OCR_TRANSLATE"], img])

        if hasattr(response, 'text') and response.text:
            translated_text = response.text.strip()
            doc = docx.Document()
            doc.add_paragraph(translated_text)
            doc.save(output_docx_path)
            logging.info(f"圖片 OCR 翻譯結果儲存成功: {os.path.basename(output_docx_path)}")
            return True
        else:
            block_reason = f" (Block Reason: {response.prompt_feedback.block_reason})" if hasattr(response, 'prompt_feedback') else ""
            raise Exception(f"Gemini API 未能生成有效的翻譯文字{block_reason}")
    except Exception as e:
        logging.error(f"處理圖片 OCR 與翻譯時發生錯誤: {e}", exc_info=True)
        if progress_queue: progress_queue.put(json.dumps({'type': 'error', 'message': f'處理圖片失敗: {e}'}))
        return False